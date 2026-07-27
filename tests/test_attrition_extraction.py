import json
import tempfile
from pathlib import Path

from docx import Document

from ads_automation.notebook_generator import create_notebook
from ads_automation.parser import parse_protocol
from app import process_protocol


def test_parse_protocol_extracts_title_and_attrition_steps():
    with tempfile.TemporaryDirectory() as tmpdir:
        path = Path(tmpdir) / "protocol.docx"
        doc = Document()
        doc.add_paragraph("Study Protocol")
        doc.add_paragraph("Research Project Title: Example Study")
        doc.add_paragraph("Data Sources: Premier Healthcare Database")
        doc.add_paragraph("Study Design")
        doc.add_paragraph("Inclusion Criteria")
        doc.add_paragraph("1. Patients aged 18 years or older")
        doc.add_paragraph("2. Patients with at least one procedure")
        doc.add_paragraph("Exclusion Criteria")
        doc.add_paragraph("1. Patients with prior surgery")
        doc.save(path)

        result = parse_protocol(path)

        assert result["title"] == "Example Study"
        assert result["data_sources"] == ["Premier Healthcare Database"]
        assert result["inclusion_steps"] == ["Patients aged 18 years or older", "Patients with at least one procedure"]
        assert result["exclusion_steps"] == ["Patients with prior surgery"]


def test_parse_protocol_handles_title_from_study_protocol_heading():
    with tempfile.TemporaryDirectory() as tmpdir:
        path = Path(tmpdir) / "protocol.docx"
        doc = Document()
        doc.add_paragraph("Study Protocol")
        doc.add_paragraph("Example Study for Attrition")
        doc.add_paragraph("Data Source(s)")
        doc.add_paragraph("Premier Healthcare Database")
        doc.add_paragraph("Study Design")
        doc.add_paragraph("Inclusion Criteria")
        doc.add_paragraph("1. Patients aged 18 years or older")
        doc.add_paragraph("2. Patients with at least one procedure")
        doc.add_paragraph("Exclusion Criteria")
        doc.add_paragraph("1. Patients with prior surgery")
        doc.save(path)

        result = parse_protocol(path)

        assert result["title"] == "Example Study for Attrition"
        assert result["data_sources"] == ["Premier Healthcare Database"]
        assert result["inclusion_steps"][0].startswith("Patients aged")
        assert result["exclusion_steps"][0].startswith("Patients with prior surgery")


def test_process_protocol_labels_exclusion_steps_correctly_in_notebook():
    with tempfile.TemporaryDirectory() as tmpdir:
        path = Path(tmpdir) / "protocol.docx"
        output_dir = Path(tmpdir) / "notebooks"
        doc = Document()
        doc.add_paragraph("Study Protocol")
        doc.add_paragraph("Research Project Title: Example Study")
        doc.add_paragraph("Data Sources: Premier Healthcare Database")
        doc.add_paragraph("Study Design")
        doc.add_paragraph("Inclusion Criteria")
        doc.add_paragraph("1. Patients aged 18 years or older")
        doc.add_paragraph("Exclusion Criteria")
        doc.add_paragraph("1. Patients with prior surgery")
        doc.save(path)

        process_protocol(path, output_dir)

        notebook_path = next(output_dir.glob("*.ipynb"))
        notebook = json.loads(notebook_path.read_text(encoding="utf-8"))
        rendered_steps = "".join(notebook["cells"][-1]["source"])

        assert "(Inclusion) Patients aged 18 years or older" in rendered_steps
        assert "(Exclusion) Patients with prior surgery" in rendered_steps


def test_create_notebook_writes_expected_cells():
    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = Path(tmpdir) / "study_notebook.ipynb"
        steps = ["First step", "Second step"]

        create_notebook(output_path, "My Study", steps)

        assert output_path.exists()
        content = json.loads(output_path.read_text(encoding="utf-8"))
        assert len(content["cells"]) == 4
        assert "My Study" in "".join(content["cells"][0]["source"])
        assert "ATTRITION STEPS" in "".join(content["cells"][2]["source"])
        assert "1. First step" in "".join(content["cells"][3]["source"])
