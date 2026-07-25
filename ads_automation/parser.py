import re
from pathlib import Path

from docx import Document


DATA_SOURCE_MASTER = {
    # PREMIER
    "premier healthcare database": "Premier Healthcare Database",
    "pinc ai": "Premier Healthcare Database",
    "pinc" : "Premier Healthcare Database",

    # OPTUM
    "optum clinformatics date of death": "Optum DOD/SES",
    "optum clinformatics": "Optum",
    "optum dod": "Optum DOD/SES",
    "optum ses": "Optum DOD/SES",

    # MARKETSCAN
    "ibm marketscan": "IBM Marketscan",
    "ccae": "IBM Marketscan",
    "mdcr": "IBM Marketscan",
    "mdcd": "IBM Marketscan",

    # OTHERS
    "mercy": "Mercy",
    "truveta": "Truveta",
    "flatiron": "Flatiron",
    "concert ai": "Concert AI",
    "komodo": "Komodo",
    "cprd": "CPRD",
    "hes": "HES",
    "salford": "Salford Royal",
    "jmdc": "JMDC",
    "loopback": "Loopback",
    "integra": "Integra",
    "healthverity": "HealthVerity",
    "connect": "Connect",
}


# ---------------------------------------------------------------------------
# Document reading
# ---------------------------------------------------------------------------

def read_docx(file):
    """Read text from a .docx file (paragraphs + table cells).

    Paragraphs come first (body order), then table cells are appended so
    that the title extractor can find titles stored in Word tables.
    Uses para.text / cell.text (python-docx properties) so tracked-change
    markup is not double-counted.
    """
    path = Path(file)
    if not path.exists():
        raise FileNotFoundError(f"Protocol file not found: {path}")

    doc = Document(str(path))
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    return "\n".join(lines), lines


# ---------------------------------------------------------------------------
# Title extraction
# ---------------------------------------------------------------------------

TITLE_LABELS = [
    "research project title",
    "project title",
    "project name",
    "study title",
    "protocol title",
    "study name",
]

_TITLE_BLACKLIST = [
    "attention",
    "confidentiality",
    "how to use",
    "list of abbreviations",
    "table of contents",
    "project background",
    "research methods",
    "research questions",
    "protocol synopsis",
]


def extract_project_title(lines):
    # 1. Look for a labelled field like "Study Title: <value>"
    for i, line in enumerate(lines):
        lower = line.lower()
        for label in TITLE_LABELS:
            if label not in lower:
                continue
            match = re.search(
                rf"{re.escape(label)}\s*:?\s*(.+)", line, flags=re.IGNORECASE
            )
            if match:
                title = match.group(1).strip()
                if title and title != ":":
                    return title
            # Label is on its own line — grab the next non-empty line
            for j in range(i + 1, min(i + 5, len(lines))):
                candidate = lines[j].strip()
                if candidate:
                    return candidate

    # 2. Title may appear just before "Study Protocol" heading
    for i, line in enumerate(lines):
        if line.strip().lower() == "study protocol":
            for j in range(i - 1, -1, -1):
                candidate = lines[j].strip()
                if candidate and len(candidate) > 10 and "attention" not in candidate.lower():
                    return candidate

    # 3. Fall back to first long, non-blacklisted line in the document header
    for line in lines[:25]:
        candidate = line.strip()
        if len(candidate) < 15:
            continue
        if any(word in candidate.lower() for word in _TITLE_BLACKLIST):
            continue
        return candidate

    return "Unknown Study"


# ---------------------------------------------------------------------------
# Study-section extraction
# ---------------------------------------------------------------------------

def extract_study_selection(text):
    """Return the slice of *text* that covers the study design / criteria block.

    End-keyword search is anchored *after* the exclusion criteria heading so
    that keywords like 'covariates' appearing between inclusion and exclusion
    sections do not truncate the result prematurely.
    """
    text_lower = text.lower()

    start_keywords = ["study design", "study population", "inclusion criteria"]
    end_keywords = [
        "exposure variable",
        "primary independent variable",
        "covariates",
        "study outcomes",
        "product codes",
    ]

    start_idx = None
    for kw in start_keywords:
        match = re.search(kw, text_lower)
        if match:
            start_idx = match.start()
            break

    if start_idx is None:
        return text

    # Anchor the end search to after the exclusion criteria block
    exc_match = re.search(r"exclusion criteria", text_lower[start_idx:])
    search_from = start_idx + exc_match.end() if exc_match else start_idx

    end_idx = len(text)
    for kw in end_keywords:
        match = re.search(kw, text_lower[search_from:])
        if match:
            end_idx = search_from + match.start()
            break

    return text[start_idx:end_idx]


# ---------------------------------------------------------------------------
# Inclusion / exclusion split
# ---------------------------------------------------------------------------

def split_criteria_sections(text):
    text_lower = text.lower()

    if "table of contents" in text_lower:
        text = text.split("table of contents")[-1]
        text_lower = text.lower()           # keep in sync after the split

    inc_matches = list(re.finditer(r"inclusion criteria", text_lower))
    exc_matches = list(re.finditer(r"exclusion criteria", text_lower))

    # Use the last occurrence to skip any TOC / summary mentions
    inc_start = inc_matches[-1].start() if inc_matches else -1
    exc_start = exc_matches[-1].start() if exc_matches else -1

    if inc_start == -1:
        inclusion_text = ""
    else:
        inclusion_text = text[inc_start: exc_start if exc_start != -1 else len(text)]

    exclusion_text = text[exc_start:] if exc_start != -1 else ""

    return inclusion_text, exclusion_text


# ---------------------------------------------------------------------------
# Step extraction
# ---------------------------------------------------------------------------

def extract_steps(section_text):
    steps = []

    for step in section_text.split("\n"):
        step = step.strip()
        step_lower = step.lower()

        if len(step) < 5:
            continue
        if "inclusion criteria" in step_lower:
            continue
        if "exclusion criteria" in step_lower:
            continue
        if "patients will be included" in step_lower:
            continue
        if "patients will be excluded" in step_lower:
            continue
        if "table of contents" in step_lower:
            continue
        if step_lower.startswith("individuals") or step_lower.startswith("see") or step_lower.startswith("product codes"):
            continue
        if " must meet all the following" in step_lower:
            continue
        if " meeting any of the following" in step_lower:
            continue

        # Strip leading bullets and numbering
        step = re.sub(r"^[•▪◦*\-]\s*", "", step)
        step = re.sub(r"^\d+\.\s*", "", step)

        steps.append(step)

    return steps


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def parse_protocol(file):
    text, lines = read_docx(file)

    title = extract_project_title(lines)
    data_sources = detect_data_source(text)

    study_section = extract_study_selection(text)
    inc_text, exc_text = split_criteria_sections(study_section)

    inc_steps = extract_steps(inc_text)
    exc_steps = extract_steps(exc_text)

    attrition = []
    step_no = 1
    for step in inc_steps:
        attrition.append((step_no, "inclusion", step))
        step_no += 1
    for step in exc_steps:
        attrition.append((step_no, "exclusion", step))
        step_no += 1

    return {
        "title": title,
        "inclusion_steps": inc_steps,
        "exclusion_steps": exc_steps,
        "attrition": attrition,
        "data_sources": data_sources,
    }


# ---------------------------------------------------------------------------
# Data source detection
# ---------------------------------------------------------------------------

def extract_data_source_section(text):
    text_lower = text.lower()

    start = re.search(r"data sources?", text_lower)
    if not start:
        return ""

    start_idx = start.start()
    end_patterns = ["study design", "study population", "endpoints", "data analyses"]
    end_idx = len(text)

    for pattern in end_patterns:
        match = re.search(pattern, text_lower[start_idx:])
        if match:
            end_idx = start_idx + match.start()
            break

    return text[start_idx:end_idx]


def _normalize_text(text):
    """Strip trademark/registered/copyright symbols and their text equivalents."""
    text = text.replace("™", " ").replace("®", " ").replace("©", " ")
    # superscript TM or R that python-docx merges directly onto words
    text = re.sub(r"\bTM\b", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<=[a-zA-Z])TM(?=\s|$)", " ", text)
    return re.sub(r"\s+", " ", text)


def detect_data_source(text):
    section = extract_data_source_section(text)
    raw = section if section else text
    search_text = _normalize_text(raw).lower()

    detected = []
    sorted_keys = sorted(DATA_SOURCE_MASTER.keys(), key=len, reverse=True)

    for key in sorted_keys:
        pattern = r"\b" + re.escape(key) + r"\b"
        if re.search(pattern, search_text):
            detected.append(DATA_SOURCE_MASTER[key])
            search_text = re.sub(pattern, "", search_text)

    return list(set(detected))
